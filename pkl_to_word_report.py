#!/usr/bin/env python3
import os
os.environ['CUDA_VISIBLE_DEVICES'] = ''
import sys
from pathlib import Path
from datetime import datetime
import tempfile

try:
    import joblib
except ImportError:
    print('Installing joblib...')
    os.system(f'{sys.executable} -m pip install joblib -q')
    import joblib

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print('Installing python-docx...')
    os.system(f'{sys.executable} -m pip install python-docx -q')
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import seaborn as sns
except ImportError:
    print('Installing matplotlib and seaborn...')
    os.system(f'{sys.executable} -m pip install matplotlib seaborn -q')
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import seaborn as sns

try:
    import numpy as np
    from sklearn.metrics import confusion_matrix, roc_curve, roc_auc_score
except ImportError:
    print('Installing scikit-learn and numpy...')
    os.system(f'{sys.executable} -m pip install numpy scikit-learn -q')
    import numpy as np
    from sklearn.metrics import confusion_matrix, roc_curve, roc_auc_score

ROOT = Path('/home/akachat/tf_env/Stacking_CNN_Melanoma')
TARGET_DIRS = [
    ROOT / 'Models' / 'Youden_LogisticRegression_E20_T0.5', 
    ROOT / 'Models' / 'Sens95_LogisticRegression_E20_T0.5',
    ROOT / 'Models' / 'Sens95_XGBoost_E20_T0.5',
    ROOT / 'Models' / 'Sens95_XGBoost_E40_T0.5',
]

OUTPUT_DIR = ROOT / 'reports'
OUTPUT_DIR.mkdir(exist_ok=True, parents=True)

METRIC_ORDER = ['Accuracy', 'Sensitivity', 'Specificity', 'Precision', 'F1', 'AUC']


def safe_load_pkl(path):
    try:
        return joblib.load(path)
    except Exception as exc:
        return {'__load_error__': str(exc)}


def normalize_metric_value(value):
    if isinstance(value, (int, float)):
        return float(value)
    return None


def format_percentage(value):
    if value is None:
        return 'N/A'
    try:
        v = float(value)
    except (TypeError, ValueError):
        return str(value)
    return f"{v*100:.2f}%" if 0.0 <= v <= 1.0 else f"{v:.2f}%"


def shade_row(row, fill='D9E1F2'):
    for cell in row:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)


def build_metrics_table(doc, model_name, metrics):
    doc.add_heading(f'Model metrics: {model_name}', level=2)
    table = doc.add_table(rows=1, cols=len(METRIC_ORDER) + 1)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    for idx, metric in enumerate(METRIC_ORDER, 1):
        hdr_cells[idx].text = metric
        hdr_cells[idx].paragraphs[0].runs[0].bold = True

    row_cells = table.add_row().cells
    row_cells[0].text = model_name
    for idx, metric in enumerate(METRIC_ORDER, 1):
        value = metrics.get(metric, metrics.get(metric.lower(), None))
        row_cells[idx].text = format_percentage(value)
        row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    return table


def add_image(doc, image_path, caption=None, width=Inches(6)):
    if image_path.exists():
        doc.add_picture(str(image_path), width=width)
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
        doc.add_paragraph()


def plot_metrics_bar(metrics, output_path, title='Metrics'):
    names = []
    values = []
    for metric in METRIC_ORDER:
        if metric in metrics:
            val = float(metrics[metric])
            names.append(metric)
            values.append(val if val > 1.0 else val * 100.0)
        elif metric.lower() in metrics:
            val = float(metrics[metric.lower()])
            names.append(metric)
            values.append(val if val > 1.0 else val * 100.0)
    if not names:
        return False
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.bar(names, values, color='steelblue')
    ax.set_title(title)
    ax.set_ylabel('Score (%)')
    ax.set_ylim(0, max(values) * 1.2 if values else 100)
    max_val = max(values) if values else 0.0
    for i, v in enumerate(values):
        ax.text(i, v + (max_val * 0.03 if max_val else 1.0), f'{v:.2f}%', ha='center')
    plt.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return True


def plot_confusion_matrix(y_true, y_pred, output_path, title='Confusion Matrix'):
    cm = confusion_matrix(y_true, y_pred)
    cm_pct = cm.astype(float) / cm.sum() * 100 if cm.sum() else np.zeros_like(cm, dtype=float)
    labels = np.array([[f'{cm[i, j]}\n({cm_pct[i, j]:.1f}%)' for j in range(cm.shape[1])] for i in range(cm.shape[0])])
    fig, ax = plt.subplots(figsize=(6, 5))
    sns.heatmap(cm, annot=labels, fmt='', cmap='Blues', ax=ax, cbar=False,
                xticklabels=['Benign', 'Malignant'], yticklabels=['Benign', 'Malignant'])
    ax.set_xlabel('Predicted')
    ax.set_ylabel('True')
    ax.set_title(title)
    plt.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return True


def plot_base_models_bar(base_metrics, output_path, title='Base Models Comparison'):
    # base_metrics: dict of model_name -> {metric: value}
    if not base_metrics:
        return False
    # collect metric names
    metric_names = set()
    for m in base_metrics.values():
        if isinstance(m, dict):
            metric_names.update(m.keys())
    # prefer METRIC_ORDER
    ordered_metrics = [mn for mn in METRIC_ORDER if mn in metric_names]
    if not ordered_metrics:
        ordered_metrics = sorted(metric_names)
    models = list(base_metrics.keys())
    if not models:
        return False
    # build data matrix
    data = np.zeros((len(ordered_metrics), len(models)))
    for j, model in enumerate(models):
        m = base_metrics.get(model, {})
        for i, metric in enumerate(ordered_metrics):
            v = m.get(metric, m.get(metric.lower(), None))
            data[i, j] = float(v) if isinstance(v, (int, float)) else 0.0
    # plot grouped bars: metrics groups along x, different colors per model
    x = np.arange(len(ordered_metrics))
    width = 0.8 / max(1, len(models))
    fig, ax = plt.subplots(figsize=(10, 5))
    for j in range(len(models)):
        ax.bar(x + j * width, data[:, j], width=width, label=models[j])
    ax.set_xticks(x + width * (len(models)-1) / 2)
    ax.set_xticklabels(ordered_metrics, rotation=45, ha='right')
    ax.set_ylabel('Score')
    ax.set_title(title)
    ax.legend()
    plt.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return True


def plot_roc_curve(y_true, y_score, output_path, title='ROC Curve'):
    if len(np.unique(y_true)) < 2:
        return False
    fpr, tpr, _ = roc_curve(y_true, y_score)
    auc = roc_auc_score(y_true, y_score)
    fig, ax = plt.subplots(figsize=(6, 5))
    ax.plot(fpr, tpr, lw=2, label=f'AUC = {auc:.3f}')
    ax.plot([0, 1], [0, 1], linestyle='--', color='gray')
    ax.set_xlabel('False Positive Rate')
    ax.set_ylabel('True Positive Rate')
    ax.set_title(title)
    ax.legend(loc='lower right')
    plt.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return True


def plot_training_history(histories, output_path, title='Training History'):
    if not histories:
        return False
    fig, axes = plt.subplots(2, 1, figsize=(8, 10), sharex=True)
    for model_name, history in histories.items():
        if not isinstance(history, dict):
            continue
        epochs = range(1, len(history.get('accuracy', [])) + 1)
        if history.get('accuracy'):
            axes[0].plot(epochs, history['accuracy'], label=f'{model_name} train')
        if history.get('val_accuracy'):
            axes[0].plot(epochs, history['val_accuracy'], linestyle='--', label=f'{model_name} val')
        if history.get('loss'):
            axes[1].plot(epochs, history['loss'], label=f'{model_name} train')
        if history.get('val_loss'):
            axes[1].plot(epochs, history['val_loss'], linestyle='--', label=f'{model_name} val')
    axes[0].set_title('Accuracy')
    axes[0].set_ylabel('Score')
    axes[0].legend(fontsize='small')
    axes[1].set_title('Loss')
    axes[1].set_xlabel('Epoch')
    axes[1].set_ylabel('Loss')
    axes[1].legend(fontsize='small')
    fig.suptitle(title)
    plt.tight_layout(rect=[0, 0.03, 1, 0.95])
    fig.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return True


def get_final_scores(data):
    if 'meta_learner' not in data or 'scaler' not in data or 'base_model_predictions' not in data:
        return None
    try:
        test_preds = []
        for model_name, preds in data['base_model_predictions'].items():
            if isinstance(preds, dict) and 'test' in preds:
                arr = np.asarray(preds['test']).ravel()
                test_preds.append(arr)
        if not test_preds:
            return None
        X_test = np.column_stack(test_preds)
        X_test = data['scaler'].transform(X_test)
        probs = data['meta_learner'].predict_proba(X_test)[:, 1]
        return probs
    except Exception:
        return None


def extract_model_data(pkl_path):
    data = safe_load_pkl(pkl_path)
    if isinstance(data, dict) and '__load_error__' in data:
        return {'error': data['__load_error__']}

    result = {'name': pkl_path.stem, 'path': str(pkl_path), 'directory': str(pkl_path.parent)}
    metrics = {}
    if isinstance(data, dict):
        raw_metrics = data.get('metrics', {})
        if isinstance(raw_metrics, dict):
            metrics = {k: normalize_metric_value(v) for k, v in raw_metrics.items() if isinstance(v, (int, float))}

        # extract base model metrics and flatten into metrics with prefixes
        base_model_metrics = data.get('base_model_metrics', {}) or {}
        result['base_model_metrics'] = base_model_metrics
        preferred_models = ['EfficientNetV2B0', 'MobileNetV2', 'DenseNet169']
        for model_name in preferred_models:
            bm = base_model_metrics.get(model_name) or base_model_metrics.get(model_name.replace('V2','V2B0')) or base_model_metrics.get(model_name.replace('MobileNetV2','MobileNetv2'))
            if isinstance(bm, dict):
                for mk, mv in bm.items():
                    if isinstance(mv, (int, float)):
                        metrics[f"{model_name}_{mk}"] = normalize_metric_value(mv)

        result['y_true'] = np.asarray(data.get('y_true', []), dtype=int) if data.get('y_true', None) is not None else np.array([])
        result['y_pred'] = np.asarray(data.get('y_pred', []), dtype=int) if data.get('y_pred', None) is not None else np.array([])
        result['histories'] = data.get('histories', {})
        result['final_scores'] = get_final_scores(data)
        result['status'] = 'loaded'
    else:
        result['status'] = 'unsupported'
    result['metrics'] = metrics
    return result


def create_report(results, output_file):
    doc = Document()
    doc.add_heading('Melanoma PKL Metrics Report', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Source directories: {", ".join(str(d) for d in TARGET_DIRS)}')
    doc.add_paragraph(f'Model files processed: {len(results)}')
    doc.add_page_break()

    # ====================================================================
    # GLOBAL METRICS TABLE (page 2) - split by source directory
    # ====================================================================
    from pathlib import Path as _Path

    groups = {}
    for r in results:
        if isinstance(r, dict):
            dir_name = _Path(r.get('directory', '')).name or 'Unknown'
        else:
            dir_name = 'Unknown'
        groups.setdefault(dir_name, []).append(r)

    # union of metric keys across all groups
    all_metric_keys = set()
    for lst in groups.values():
        for r in lst:
            if isinstance(r, dict):
                m = r.get('metrics', {})
                if isinstance(m, dict):
                    all_metric_keys.update(m.keys())

    preferred = ['Accuracy', 'Sensitivity', 'Specificity', 'Precision', 'F1', 'AUC']
    metric_cols = [m for m in preferred if m in all_metric_keys]
    metric_cols.extend([m for m in sorted(all_metric_keys) if m not in metric_cols])

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading('Global Metrics Table (grouped)', level=1)

    if not metric_cols:
        doc.add_paragraph('No numeric metrics found across processed PKL files.')
    else:
        metric_names = ['Accuracy', 'Sensitivity', 'Specificity', 'Precision', 'F1', 'AUC']
        table = doc.add_table(rows=1, cols=3 + len(metric_names))
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Group'
        hdr[1].text = 'PKL'
        hdr[2].text = 'Model'
        for i, metric in enumerate(metric_names, start=3):
            hdr[i].text = metric

        def format_value(value):
            return format_percentage(value)

        for group_name in sorted(groups.keys()):
            for r in groups[group_name]:
                base_metrics = r.get('base_model_metrics', {}) if isinstance(r, dict) else {}
                meta_metrics = r.get('metrics', {}) if isinstance(r, dict) else {}

                # Meta learner row
                row = table.add_row().cells
                row[0].text = group_name
                row[1].text = r.get('name', 'Unknown PKL')
                row[2].text = 'Meta learner'
                for col_idx, metric in enumerate(metric_names, start=3):
                    val = meta_metrics.get(metric, meta_metrics.get(metric.lower(), None))
                    row[col_idx].text = format_value(val)
                shade_row(row)

                # Base model rows
                for model_name in ['EfficientNetV2B0', 'MobileNetV2', 'DenseNet169']:
                    row = table.add_row().cells
                    row[0].text = group_name
                    row[1].text = r.get('name', 'Unknown PKL')
                    row[2].text = model_name
                    bm = base_metrics.get(model_name, {}) if isinstance(base_metrics, dict) else {}
                    if not isinstance(bm, dict):
                        bm = {}
                    for col_idx, metric in enumerate(metric_names, start=3):
                        val = bm.get(metric, bm.get(metric.lower(), None))
                        row[col_idx].text = format_value(val)
                    shade_row(row, fill='FFFFFF')

    doc.add_page_break()
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width

    for res in results:
        doc.add_heading(f"{res['directory'].split('/')[-1]} / {res['name']}", level=1)
        if res.get('error'):
            doc.add_paragraph(f"Error loading PKL: {res['error']}")
            doc.add_page_break()
            continue
        build_metrics_table(doc, res['name'], res['metrics'])

        # Per-base-model metrics table and placeholder for bar plot
        base_metrics = res.get('base_model_metrics', {}) if isinstance(res, dict) else {}
        base_present = bool(base_metrics)
        if base_present:
            doc.add_heading('Per-base-model metrics', level=3)
            # collect metric keys across base models
            base_metric_keys = set()
            for bm in base_metrics.values():
                if isinstance(bm, dict):
                    base_metric_keys.update(bm.keys())
            base_metric_order = [m for m in METRIC_ORDER if m in base_metric_keys]
            base_metric_order.extend([m for m in sorted(base_metric_keys) if m not in base_metric_order])
            if not base_metric_order:
                doc.add_paragraph('No numeric metrics found for base models.')
            else:
                rows = 1 + len(base_metrics)
                cols = 1 + len(base_metric_order)
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Light Grid Accent 1'
                hdr = table.rows[0].cells
                hdr[0].text = 'Model'
                for i, col in enumerate(base_metric_order, start=1):
                    hdr[i].text = col
                for row_idx, model_name in enumerate(sorted(base_metrics.keys()), start=1):
                    row = table.rows[row_idx].cells
                    row[0].text = model_name
                    bm = base_metrics.get(model_name, {})
                    for col_idx, col in enumerate(base_metric_order, start=1):
                        val = bm.get(col, bm.get(col.lower(), None))
                        row[col_idx].text = format_percentage(val)
                doc.add_paragraph()

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            bar_path = tmpdir / f"{res['name']}_metrics.png"
            cm_path = tmpdir / f"{res['name']}_cm.png"
            roc_path = tmpdir / f"{res['name']}_roc.png"
            history_path = tmpdir / f"{res['name']}_history.png"
            base_bar_path = tmpdir / f"{res['name']}_base_metrics.png"

            if plot_metrics_bar(res['metrics'], bar_path, title=f"Metrics for {res['name']}"):
                add_image(doc, bar_path, caption='Bar plot des métriques')

            # base models bar plot
            if base_present:
                try:
                    if plot_base_models_bar(base_metrics, base_bar_path, title=f"Base models - {res['name']}"):
                        add_image(doc, base_bar_path, caption='Comparaison des 3 modèles de base')
                except Exception:
                    pass

            if res['y_true'].size and res['y_pred'].size:
                if plot_confusion_matrix(res['y_true'], res['y_pred'], cm_path, title='Confusion Matrix'):
                    add_image(doc, cm_path, caption='Matrice de confusion')
            if res['final_scores'] is not None and res['y_true'].size:
                if plot_roc_curve(res['y_true'], res['final_scores'], roc_path, title='ROC Curve'):
                    add_image(doc, roc_path, caption='AUC_ROC curve')
            elif 'AUC' in res['metrics']:
                doc.add_paragraph(f"AUC score: {res['metrics']['AUC']:.2f}")

            if plot_training_history(res.get('histories', {}), history_path, title='Training History'):
                add_image(doc, history_path, caption='PLOT train historique')

        doc.add_page_break()

    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)
    return output_file


def main():
    pkl_files = []
    for folder in TARGET_DIRS:
        if folder.exists() and folder.is_dir():
            pkl_files.extend(sorted(folder.glob('*.pkl')))
        else:
            print(f'Warning: target folder not found: {folder}')
    if not pkl_files:
        print('No pickle files found to process.')
        return 1

    print(f'Found {len(pkl_files)} pickle files.')
    results = []
    for pkl_path in pkl_files:
        print(f'Processing {pkl_path.name}...')
        results.append(extract_model_data(pkl_path))

    output_file = OUTPUT_DIR / f'PKL_Report_Indice_de_Youden_Sensibilite_sup_95_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    report_path = create_report(results, output_file)
    print(f'Report saved to: {report_path}')
    return 0

if __name__ == '__main__':
    sys.exit(main())
